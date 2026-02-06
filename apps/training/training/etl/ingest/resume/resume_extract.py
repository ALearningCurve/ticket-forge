"""Step 1: Resume Extractor - Extracts raw text from PDF/DOCX/DOC files."""

import os
import uuid
from pathlib import Path
from typing import List
from dataclasses import dataclass

import fitz
from docx import Document

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False


@dataclass
class ExtractedResume:
    engineer_id: str
    filename: str
    raw_content: str


class ResumeExtractor:
    SUPPORTED = {'.pdf', '.docx', '.doc'}

    def __init__(self, id_prefix: str = "ENG", use_uuid: bool = False):
        self.id_prefix = id_prefix
        self.use_uuid = use_uuid
        self._counter = 0

    def _generate_id(self) -> str:
        if self.use_uuid:
            return str(uuid.uuid4())[:8].upper()
        self._counter += 1
        return f"{self.id_prefix}-{self._counter:03d}"

    def extract(self, file_path: str) -> str:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.SUPPORTED}")

        if ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext == '.docx':
            return self._extract_docx(file_path)
        elif ext == '.doc':
            return self._extract_doc(file_path)

    def _extract_pdf(self, path: str) -> str:
        text_parts = []
        with fitz.open(path) as doc:
            for page in doc:
                text_parts.append(page.get_text())

        text = "\n".join(text_parts).strip()

        if len(text) < 100 and HAS_OCR:
            return self._extract_pdf_ocr(path)

        return text

    def _extract_pdf_ocr(self, path: str) -> str:
        text_parts = []
        with fitz.open(path) as doc:
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text_parts.append(pytesseract.image_to_string(img))

        return "\n".join(text_parts).strip()

    def _extract_docx(self, path: str) -> str:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)

        return "\n".join(parts).strip()

    def _extract_doc(self, path: str) -> str:
        if not HAS_WIN32:
            raise RuntimeError("DOC extraction requires pywin32 on Windows.")

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(path))
            text = doc.Content.Text
            doc.Close()
            return text.strip()
        finally:
            word.Quit()

    def extract_directory(self, directory: str) -> List[ExtractedResume]:
        dir_path = Path(directory)
        results = []
        files = [f for f in dir_path.iterdir() if f.suffix.lower() in self.SUPPORTED]

        if not files:
            return results

        for file_path in files:
            try:
                text = self.extract(str(file_path))
                results.append(ExtractedResume(
                    engineer_id=self._generate_id(),
                    filename=file_path.name,
                    raw_content=text
                ))
            except Exception:
                continue

        return results