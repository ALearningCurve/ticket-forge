"""Resume upload and ML profile update service for authenticated users.

This module provides a direct web-backend integration for resume upload so the
profile flow can run inside the deployed API service without requiring Airflow.
It reuses shared ML components from ``ml-core`` for embeddings, keywords, and
vector formatting, then persists results into the existing pgvector-backed
``users`` table and bridges that record to ``auth_users``.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path

import fitz
from docx import Document
from ml_core.embeddings import get_embedding_service
from ml_core.keywords import get_keyword_extractor
from ml_core.profiles import ProfileUpdater
from ml_core.retrieval.hybrid_retrieval import vector_to_pgvector_text
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from web_backend.models.user import AuthUser

MAX_RESUME_SIZE_BYTES = 10 * 1024 * 1024
SIMILARITY_THRESHOLD = 0.01
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

_WHITESPACE_RE = re.compile(r"\s+")
_PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}")
_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
_URL_RE = re.compile(r"(?:https?://|www\.)\S+")


@dataclass(slots=True)
class ResumeProfileStatus:
    """Resume/profile status for an authenticated user."""

    has_resume: bool
    member_id: int | None
    last_uploaded_at: datetime | None


@dataclass(slots=True)
class ResumeProfileUploadResult(ResumeProfileStatus):
    """Resume upload result for an authenticated user."""

    action: str


def _normalize_resume_text(text_value: str) -> str:
    """Remove obvious PII and normalize whitespace from resume text.

    Args:
        text_value: Extracted raw resume text.

    Returns:
        Cleaned text suitable for embeddings and keyword extraction.
    """
    normalized = _PHONE_RE.sub(" ", text_value)
    normalized = _EMAIL_RE.sub(" ", normalized)
    normalized = _URL_RE.sub(" ", normalized)
    normalized = normalized.replace("\x00", " ")
    lines = [line.strip() for line in normalized.splitlines() if line.strip()]
    return _WHITESPACE_RE.sub(" ", "\n".join(lines)).strip()


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a PDF upload.

    Args:
        file_bytes: Raw PDF bytes.

    Returns:
        Extracted text content.

    Raises:
        ValueError: If no meaningful text can be extracted.
    """
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        pages = [page.get_text("text") for page in document]
    extracted = "\n".join(pages).strip()
    if extracted:
        return extracted
    msg = "Could not extract text from the uploaded PDF."
    raise ValueError(msg)


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a DOCX upload.

    Args:
        file_bytes: Raw DOCX bytes.

    Returns:
        Extracted text content.

    Raises:
        ValueError: If no meaningful text can be extracted.
    """
    document = Document(BytesIO(file_bytes))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    extracted = "\n".join(parts).strip()
    if extracted:
        return extracted
    msg = "Could not extract text from the uploaded DOCX."
    raise ValueError(msg)


def _extract_resume_text(filename: str, file_bytes: bytes) -> str:
    """Extract raw text from a supported resume upload.

    Args:
        filename: Uploaded filename.
        file_bytes: Uploaded file bytes.

    Returns:
        Extracted raw text.

    Raises:
        ValueError: If the file is too large or the extension is unsupported.
    """
    if not filename:
        msg = "Uploaded file must include a filename."
        raise ValueError(msg)
    if len(file_bytes) > MAX_RESUME_SIZE_BYTES:
        msg = "Resume file must be 10MB or smaller."
        raise ValueError(msg)

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        msg = "Only PDF and DOCX resume files are supported."
        raise ValueError(msg)

    if extension == ".pdf":
        return _extract_pdf_text(file_bytes)
    return _extract_docx_text(file_bytes)


async def get_resume_profile_status(
    db: AsyncSession,
    current_user: AuthUser,
) -> ResumeProfileStatus:
    """Return resume/profile status for the current authenticated user.

    Args:
        db: Active async database session.
        current_user: Authenticated application user.

    Returns:
        Status describing whether an ML profile resume is on file.
    """
    if current_user.member_id is None:
        return ResumeProfileStatus(
            has_resume=False,
            member_id=None,
            last_uploaded_at=None,
        )

    result = await db.execute(
        text(
            """
            SELECT
              member_id,
              (resume_base_vector IS NOT NULL) AS has_resume,
              updated_at AS last_uploaded_at
            FROM users
            WHERE member_id = :member_id
            """
        ),
        {"member_id": current_user.member_id},
    )
    row = result.mappings().first()
    if row is None:
        return ResumeProfileStatus(
            has_resume=False,
            member_id=current_user.member_id,
            last_uploaded_at=None,
        )
    return ResumeProfileStatus(
        has_resume=bool(row["has_resume"]),
        member_id=int(row["member_id"]),
        last_uploaded_at=row["last_uploaded_at"],
    )


async def upload_resume_for_user(
    db: AsyncSession,
    current_user: AuthUser,
    *,
    filename: str,
    file_bytes: bytes,
) -> ResumeProfileUploadResult:
    """Process a resume upload and bridge it to the authenticated user.

    Args:
        db: Active async database session.
        current_user: Authenticated application user.
        filename: Uploaded filename.
        file_bytes: Uploaded file bytes.

    Returns:
        Upload result with the resulting ML profile status.

    Raises:
        ValueError: If the upload is invalid or cannot be processed.
    """
    raw_text = _extract_resume_text(filename, file_bytes)
    normalized_text = _normalize_resume_text(raw_text)
    if not normalized_text:
        msg = "The uploaded resume did not contain usable text."
        raise ValueError(msg)

    embedding_service = get_embedding_service()
    keyword_extractor = get_keyword_extractor()
    updater = ProfileUpdater()

    raw_embedding = embedding_service.embed_text(normalized_text)
    vector_values = (
        raw_embedding.tolist() if hasattr(raw_embedding, "tolist") else raw_embedding
    )
    embedding = [float(value) for value in vector_values]
    vector_text = vector_to_pgvector_text(embedding)

    keywords = keyword_extractor.extract(normalized_text)
    keywords_text = " ".join(keywords)
    full_name = f"{current_user.first_name} {current_user.last_name}".strip()

    lookup_params = {
        "member_id": current_user.member_id,
        "github_username": current_user.username,
        "resume_vector": vector_text,
    }
    row = None
    if current_user.member_id is not None:
        by_member = await db.execute(
            text(
                """
                SELECT
                  member_id,
                  (resume_base_vector IS NULL) AS is_stub,
                  CASE
                    WHEN resume_base_vector IS NULL THEN NULL
                    ELSE (resume_base_vector <=> CAST(:resume_vector AS vector))
                  END AS cosine_dist
                FROM users
                WHERE member_id = :member_id
                """
            ),
            lookup_params,
        )
        row = by_member.mappings().first()

    if row is None:
        by_username = await db.execute(
            text(
                """
                SELECT
                  member_id,
                  (resume_base_vector IS NULL) AS is_stub,
                  CASE
                    WHEN resume_base_vector IS NULL THEN NULL
                    ELSE (resume_base_vector <=> CAST(:resume_vector AS vector))
                  END AS cosine_dist
                FROM users
                WHERE github_username = :github_username
                """
            ),
            lookup_params,
        )
        row = by_username.mappings().first()

    action = "created"
    if row is None:
        created = await db.execute(
            text(
                """
                INSERT INTO users (
                  github_username,
                  full_name,
                  resume_base_vector,
                  profile_vector,
                  skill_keywords
                )
                VALUES (
                  :github_username,
                  :full_name,
                  CAST(:resume_vector AS vector),
                  CAST(:resume_vector AS vector),
                  to_tsvector('english', :keywords_text)
                )
                RETURNING member_id, updated_at
                """
            ),
            {
                "github_username": current_user.username,
                "full_name": full_name,
                "resume_vector": vector_text,
                "keywords_text": keywords_text,
            },
        )
        result_row = created.mappings().one()
        action = "created"
    elif bool(row["is_stub"]):
        enriched = await db.execute(
            text(
                """
                UPDATE users
                SET
                  github_username = :github_username,
                  full_name = :full_name,
                  resume_base_vector = CAST(:resume_vector AS vector),
                  profile_vector = CAST(:resume_vector AS vector),
                  skill_keywords = to_tsvector('english', :keywords_text),
                  updated_at = now()
                WHERE member_id = :member_id
                RETURNING member_id, updated_at
                """
            ),
            {
                "github_username": current_user.username,
                "full_name": full_name,
                "resume_vector": vector_text,
                "keywords_text": keywords_text,
                "member_id": row["member_id"],
            },
        )
        result_row = enriched.mappings().one()
        action = "enriched"
    elif (
        row["cosine_dist"] is not None
        and float(row["cosine_dist"]) < SIMILARITY_THRESHOLD
    ):
        skipped = await db.execute(
            text(
                """
                UPDATE users
                SET
                  github_username = :github_username,
                  full_name = :full_name,
                  updated_at = now()
                WHERE member_id = :member_id
                RETURNING member_id, updated_at
                """
            ),
            {
                "github_username": current_user.username,
                "full_name": full_name,
                "member_id": row["member_id"],
            },
        )
        result_row = skipped.mappings().one()
        action = "skipped"
    else:
        updated = await db.execute(
            text(
                """
                UPDATE users
                SET
                  github_username = :github_username,
                  full_name = :full_name,
                  resume_base_vector = CAST(:resume_vector AS vector),
                  profile_vector =
                    (array_fill(CAST(:alpha AS real), ARRAY[384])::vector * profile_vector
                     + array_fill(CAST(:one_minus_alpha AS real), ARRAY[384])::vector
                     * CAST(:resume_vector AS vector)),
                  skill_keywords =
                    skill_keywords || to_tsvector('english', :keywords_text),
                  updated_at = now()
                WHERE member_id = :member_id
                RETURNING member_id, updated_at
                """
            ),
            {
                "github_username": current_user.username,
                "full_name": full_name,
                "resume_vector": vector_text,
                "keywords_text": keywords_text,
                "member_id": row["member_id"],
                "alpha": updater.alpha,
                "one_minus_alpha": 1.0 - updater.alpha,
            },
        )
        result_row = updated.mappings().one()
        action = "updated"

    member_id = int(result_row["member_id"])
    if current_user.member_id != member_id:
        current_user.member_id = member_id

    await db.commit()
    return ResumeProfileUploadResult(
        action=action,
        has_resume=True,
        member_id=member_id,
        last_uploaded_at=result_row["updated_at"],
    )
