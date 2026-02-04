"""
Step 1: Resume Extractor
------------------------
Extracts raw text from PDF/DOCX/DOC files.

Install:
    pip install pymupdf python-docx
"""

import os
import uuid
from pathlib import Path
from typing import Dict, List
from dataclasses import dataclass

import fitz  # PyMuPDF
from docx import Document

# OCR (optional)
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# Windows .doc (optional)
try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False


@dataclass
class ExtractedResume:
    """Extracted resume data with unique ID."""
    engineer_id: str      # Unique ID (e.g., ENG-001 or UUID)
    filename: str         # Original filename
    raw_content: str      # Extracted text


class ResumeExtractor:
    """Extract raw text from resume files."""

    SUPPORTED = {'.pdf', '.docx', '.doc'}

    def __init__(self, id_prefix: str = "ENG", use_uuid: bool = False):
        """
        Initialize extractor.

        Args:
            id_prefix: Prefix for sequential IDs (e.g., "ENG" -> ENG-001, ENG-002)
            use_uuid: If True, use UUID instead of sequential IDs
        """
        self.id_prefix = id_prefix
        self.use_uuid = use_uuid
        self._counter = 0

    def _generate_id(self) -> str:
        """Generate unique engineer ID."""
        if self.use_uuid:
            return str(uuid.uuid4())[:8].upper()
        else:
            self._counter += 1
            return f"{self.id_prefix}-{self._counter:03d}"

    def extract(self, file_path: str) -> str:
        """
        Extract text from a single resume file.

        Args:
            file_path: Path to PDF/DOCX/DOC file

        Returns:
            Extracted raw text
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.SUPPORTED}")

        if ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext == '.docx':
            return self._extract_docx(file_path)
        elif ext == '.doc':
            return self._extract_doc(file_path)

    def _extract_pdf(self, path: str) -> str:
        """Extract text from PDF, with OCR fallback."""
        text_parts = []

        with fitz.open(path) as doc:
            for page in doc:
                text_parts.append(page.get_text())

        text = "\n".join(text_parts).strip()

        # If minimal text, try OCR (likely scanned PDF)
        if len(text) < 100:
            if HAS_OCR:
                print(f"    → Low text detected, using OCR...")
                return self._extract_pdf_ocr(path)
            else:
                print(f"    ⚠️ Low text and OCR not available")

        return text

    def _extract_pdf_ocr(self, path: str) -> str:
        """Extract text from scanned PDF using OCR."""
        text_parts = []

        with fitz.open(path) as doc:
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text_parts.append(pytesseract.image_to_string(img))

        return "\n".join(text_parts).strip()

    def _extract_docx(self, path: str) -> str:
        """Extract text from DOCX."""
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]

        # Include table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)

        return "\n".join(parts).strip()

    def _extract_doc(self, path: str) -> str:
        """Extract text from DOC (Windows only)."""
        if not HAS_WIN32:
            raise RuntimeError(
                "DOC extraction requires pywin32 on Windows. "
                "Consider converting to DOCX."
            )

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(path))
            text = doc.Content.Text
            doc.Close()
            return text.strip()
        finally:
            word.Quit()

    def extract_directory(self, directory: str) -> List[ExtractedResume]:
        """
        Extract text from all resumes in a directory.

        Args:
            directory: Path to directory containing resume files

        Returns:
            List of ExtractedResume objects with unique IDs
        """
        dir_path = Path(directory)
        results = []

        files = [f for f in dir_path.iterdir() if f.suffix.lower() in self.SUPPORTED]

        if not files:
            print(f"No supported files found in {directory}")
            return results

        print(f"Found {len(files)} resume(s)\n")

        for file_path in files:
            engineer_id = self._generate_id()
            filename = file_path.name

            print(f"Extracting: {filename} → {engineer_id}")

            try:
                text = self.extract(str(file_path))

                extracted = ExtractedResume(
                    engineer_id=engineer_id,
                    filename=filename,
                    raw_content=text
                )
                results.append(extracted)

                print(f"    ✓ {len(text.split())} words extracted")
            except Exception as e:
                print(f"    ✗ Error: {e}")

        return results